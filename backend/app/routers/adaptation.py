"""剧本改编模块路由。"""
import asyncio
import json
import logging
import re
from app.core.datetime_utils import utcnow_naive
from io import BytesIO
from typing import List, Optional

_SCENE_TITLE_PATTERNS = [
    re.compile(r"^\d+[-－—]\d+"),                # 1-1, 2-3
    re.compile(r"^场\s*\d+"),                    # 场1
    re.compile(r"^第[一二三四五六七八九十百千\d]+场"),  # 第一场
]
_CHAR_NAME_PATTERN = re.compile(r"([一-鿿]{2,4})(?:（|Vo|OS|：)")

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select, func, delete
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.core.security import create_sse_ticket, verify_sse_ticket
from app.models.user import User
from app.models.adaptation_project import AdaptationProject
from app.models.adaptation_mapping_entry import AdaptationMappingEntry
from app.models.adaptation_version import AdaptationVersion
from app.models.adaptation_scene_result import AdaptationSceneResult
from app.routers.auth import get_current_user
from app.schemas.adaptation import (
    AdaptationProjectCreate, AdaptationProjectUpdate, AdaptationProjectOut,
    AdaptationVersionOut, MappingsBulkPut, MappingEntryOut, SceneBoundary,
    MappingSuggestRequest, RunCreate, SceneRerunRequest, SceneManualPatch,
    VersionDetailOut, SceneResultOut,
)
from app.services.adaptation_pipeline import AdaptationPipeline
from app.services.adaptation_llm_service import get_default_service, _strip_code_fence
from app.services.adaptation_event_bus import event_bus
from app.services.file_parser import FileParser

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/v1/adaptation", tags=["adaptation"])


def _word_count(text: str) -> int:
    import re
    return len(re.findall(r"[一-鿿]", text)) + len(re.findall(r"[a-zA-Z]+", text))


def _parse_file(content: bytes, filename: str) -> str:
    """按文件扩展名分发解析。"""
    fname = filename.lower()
    if fname.endswith(".docx"):
        return FileParser.parse_docx(content).text
    elif fname.endswith(".md") or fname.endswith(".markdown"):
        return FileParser.parse_markdown(content).text
    else:
        return FileParser.parse_txt(content).text


async def _get_owned_project(
    project_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> AdaptationProject:
    from sqlalchemy.orm import selectinload
    p = (await db.execute(
        select(AdaptationProject)
        .where(AdaptationProject.id == project_id)
        .options(selectinload(AdaptationProject.versions), selectinload(AdaptationProject.mappings))
    )).scalar_one_or_none()
    if not p or (p.user_id != current_user.id and not getattr(current_user, "is_superuser", False)):
        raise HTTPException(status_code=404, detail="改编项目不存在或无权访问")
    return p


def _project_to_out(p: AdaptationProject) -> dict:
    meta = p.metadata_ or {}
    return {
        "id": p.id, "title": p.title, "source_filename": p.source_filename,
        "intent": p.intent, "intensity": p.intensity, "era_target": p.era_target,
        "status": p.status, "created_at": p.created_at, "updated_at": p.updated_at,
        "word_count": _word_count(p.source_text),
        "scene_boundaries": meta.get("scene_boundaries", []),
        "versions": [
            {
                "id": v.id, "version_no": v.version_no, "triggered_by": v.triggered_by,
                "status": v.status, "stats": v.stats, "error": v.error,
                "created_at": v.created_at, "completed_at": v.completed_at,
            } for v in p.versions
        ],
        "mappings": [
            {
                "id": m.id, "entity_type": m.entity_type,
                "original_text": m.original_text, "replacement_text": m.replacement_text,
                "locked": m.locked, "notes": m.notes, "order_index": m.order_index,
            } for m in p.mappings
        ],
    }


def _version_to_out(v) -> dict:
    return {
        "id": v.id, "version_no": v.version_no, "triggered_by": v.triggered_by,
        "status": v.status, "stats": v.stats, "error": v.error,
        "created_at": v.created_at, "completed_at": v.completed_at,
    }


def _scene_to_out(s) -> dict:
    return {
        "id": s.id, "scene_index": s.scene_index, "scene_title": s.scene_title,
        "status": s.status, "error": s.error, "token_used": s.token_used,
        "line_count_delta_pct": s.line_count_delta_pct,
        "original_scene_text": s.original_scene_text,
        "rewritten_scene_text": s.rewritten_scene_text,
        "manual_edits": s.manual_edits or [],
        "updated_at": s.updated_at,
    }


# ─── Project CRUD ────────────────────────────────────────────────────────

@router.post("/projects", status_code=201, response_model=AdaptationProjectOut)
async def create_project(
    payload: AdaptationProjectCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if not payload.raw_text:
        raise HTTPException(400, "raw_text 不能为空（上传文件请用 /projects/upload 端点）")
    if len(payload.raw_text) > settings.ADAPTATION_MAX_CHARS:
        raise HTTPException(400, f"原文超过 {settings.ADAPTATION_MAX_CHARS} 字上限")
    p = AdaptationProject(
        user_id=current_user.id, title=payload.title,
        source_text=payload.raw_text, intent=payload.intent,
        intensity=payload.intensity, era_target=payload.era_target,
        status="ready", metadata_={},
    )
    db.add(p); await db.commit(); await db.refresh(p)
    # Eager load relationships to avoid lazy-load in async context
    from sqlalchemy.orm import selectinload
    await db.refresh(p, attribute_names=["versions", "mappings"])
    return _project_to_out(p)


@router.post("/projects/upload", status_code=201, response_model=AdaptationProjectOut)
async def create_project_upload(
    title: str = Form(...),
    intensity: int = Form(2),
    intent: Optional[str] = Form(None),
    era_target: Optional[str] = Form(None),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    content = await file.read()
    try:
        text = _parse_file(content, file.filename or "upload.txt")
    except Exception as e:
        raise HTTPException(400, f"文件解析失败：{e}")
    if len(text) > settings.ADAPTATION_MAX_CHARS:
        raise HTTPException(400, f"原文超过 {settings.ADAPTATION_MAX_CHARS} 字上限")
    p = AdaptationProject(
        user_id=current_user.id, title=title, source_filename=file.filename,
        source_text=text, intent=intent, intensity=intensity,
        era_target=era_target, status="ready", metadata_={},
    )
    db.add(p); await db.commit(); await db.refresh(p, attribute_names=["versions", "mappings"])
    return _project_to_out(p)


@router.get("/projects", response_model=List[AdaptationProjectOut])
async def list_projects(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    from sqlalchemy.orm import selectinload
    rows = (await db.execute(
        select(AdaptationProject)
        .where(AdaptationProject.user_id == current_user.id)
        .options(selectinload(AdaptationProject.versions), selectinload(AdaptationProject.mappings))
        .order_by(AdaptationProject.created_at.desc())
    )).scalars().all()
    return [_project_to_out(p) for p in rows]


@router.get("/projects/{project_id}", response_model=AdaptationProjectOut)
async def get_project(p: AdaptationProject = Depends(_get_owned_project)):
    return _project_to_out(p)


@router.patch("/projects/{project_id}", response_model=AdaptationProjectOut)
async def update_project(
    payload: AdaptationProjectUpdate,
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    for k, v in payload.model_dump(exclude_unset=True).items():
        setattr(p, k, v)
    await db.commit()
    # refresh 不带 attribute_names 会让 _get_owned_project 预加载的 versions /
    # mappings 关系过期，_project_to_out 再访问触发 N+1 lazy load。显式重水合。
    await db.refresh(p, attribute_names=["versions", "mappings"])
    return _project_to_out(p)


@router.delete("/projects/{project_id}", status_code=204)
async def delete_project(
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    await db.delete(p); await db.commit()


# ─── Extract / Split / Mappings ──────────────────────────────────────────

@router.post("/projects/{project_id}/extract", response_model=AdaptationProjectOut)
async def extract(
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    pipe = AdaptationPipeline(db=db, llm=get_default_service())
    try:
        await pipe.extract(p)
    except Exception as e:
        raise HTTPException(502, f"实体抽取失败：{e}")
    await db.refresh(p, attribute_names=["versions", "mappings"])
    return _project_to_out(p)


@router.post("/projects/{project_id}/split", response_model=AdaptationProjectOut)
async def split(
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    pipe = AdaptationPipeline(db=db, llm=get_default_service())
    await pipe.split(p)
    await db.refresh(p, attribute_names=["versions", "mappings"])
    return _project_to_out(p)


@router.put("/projects/{project_id}/mappings", response_model=List[MappingEntryOut])
async def put_mappings(
    payload: MappingsBulkPut,
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    await db.execute(
        delete(AdaptationMappingEntry).where(AdaptationMappingEntry.project_id == p.id)
    )
    for entry in payload.entries:
        db.add(AdaptationMappingEntry(project_id=p.id, **entry.model_dump()))
    await db.commit()
    rows = (await db.execute(
        select(AdaptationMappingEntry).where(AdaptationMappingEntry.project_id == p.id)
        .order_by(AdaptationMappingEntry.order_index)
    )).scalars().all()
    return [
        {"id": m.id, "entity_type": m.entity_type, "original_text": m.original_text,
         "replacement_text": m.replacement_text, "locked": m.locked,
         "notes": m.notes, "order_index": m.order_index}
        for m in rows
    ]


@router.post("/projects/{project_id}/mappings/suggest", response_model=List[MappingEntryOut])
async def suggest_mappings(
    payload: MappingSuggestRequest,
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    rows = (await db.execute(
        select(AdaptationMappingEntry).where(AdaptationMappingEntry.project_id == p.id)
        .order_by(AdaptationMappingEntry.order_index)
    )).scalars().all()

    targets = [r for r in rows if (not payload.only_empty) or not r.replacement_text]
    if not targets:
        return [{"id": m.id, "entity_type": m.entity_type, "original_text": m.original_text,
                 "replacement_text": m.replacement_text, "locked": m.locked,
                 "notes": m.notes, "order_index": m.order_index} for m in rows]

    svc = get_default_service()
    prompt = (
        f"你是剧本改编的「去识别化」命名师。新时代/世界设定：{p.era_target or '未指定'}；"
        f"改编意图：{p.intent or '未指定'}。\n"
        "\n"
        "【核心目标】为下列原词生成新名称，要让读者**无法第一时间联想到原作**。\n"
        "\n"
        "【强制规避规则（任一违反都视为失败）】\n"
        "1. 字面零重叠：新名禁止包含原名中的任何一个汉字（含姓、名、字、号、地名用字）。\n"
        "2. 音近规避：禁止与原名出现同音字、同声母+同韵母、或谐音变体（例：张→章/掌/仉 均禁）。\n"
        "3. 结构错位：原名两字则新名优先三字，原名三字可换为两字或四字；原名为叠字（婷婷/朵朵）禁止用叠字结构。\n"
        "4. 姓氏全换：人物姓氏必须更换为完全不同的姓；不得整批沿用原姓系列。\n"
        "5. 意译过近规避：地名/道具/时代词不得使用近义替换（长安→永安、剑→刃 这类一眼可猜的禁止）。\n"
        "\n"
        "【按类型生成指引】\n"
        "- person：保留原人物的性别气质/辈分/民族隐含特征（如果有），但用全新的字组合；姓与名都换。\n"
        "- place：按新时代/地域风格重新命名，避免照搬原地名的方位字（东/西/南/北）或地貌字（山/水/江/河）若与原名相同。\n"
        "- prop：按新时代背景给出功能相当但名称迥异的物件（如古剑→战术匕首；玉佩→定制胸针）。\n"
        "- era_term：用新时代背景下的对应术语，但不沿用原词任何汉字。\n"
        "\n"
        "【输出格式】严格 JSON，禁止任何解释或前后缀：\n"
        "[{\"original\": \"...\", \"replacement\": \"...\"}]\n"
        "\n"
        "【待处理原词】\n"
        + json.dumps([{"original": t.original_text, "type": t.entity_type} for t in targets], ensure_ascii=False)
    )
    raw = await svc.provider.complete(prompt, model=svc.extract_model)
    sugg: dict[str, str] = {}
    try:
        cleaned = _strip_code_fence(raw)
        sugg = {item["original"]: item["replacement"] for item in json.loads(cleaned)}
    except Exception as e:
        logger.warning("suggest_mappings 解析失败：%s; raw=%s", e, (raw or "")[:200])
    overlapped: list[tuple[str, str, str]] = []
    for t in targets:
        new_name = (sugg.get(t.original_text) or "").strip()
        if not new_name or t.locked:
            continue
        if new_name == t.original_text:
            # 完全相同毫无替换价值，跳过让用户手动处理
            continue
        overlap = set(t.original_text) & set(new_name)
        if overlap:
            # 仅记录提示，不丢弃；让用户在 UI 上自行调整或锁定
            overlapped.append((t.original_text, new_name, "".join(sorted(overlap))))
        t.replacement_text = new_name
    if overlapped:
        logger.info(
            "suggest_mappings 保留 %d 条字面重叠建议（请人工复核）：%s",
            len(overlapped),
            "; ".join(f"{o}→{n}(重叠:{ov})" for o, n, ov in overlapped),
        )
    await db.commit()
    rows = (await db.execute(
        select(AdaptationMappingEntry).where(AdaptationMappingEntry.project_id == p.id)
        .order_by(AdaptationMappingEntry.order_index)
    )).scalars().all()
    return [{"id": m.id, "entity_type": m.entity_type, "original_text": m.original_text,
             "replacement_text": m.replacement_text, "locked": m.locked,
             "notes": m.notes, "order_index": m.order_index} for m in rows]


# ─── Runs / SSE / Rerun / Patch / Export ─────────────────────────────────

async def _get_owned_version(
    version_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> AdaptationVersion:
    v = (await db.execute(
        select(AdaptationVersion).where(AdaptationVersion.id == version_id)
    )).scalar_one_or_none()
    if not v:
        raise HTTPException(404, "version 不存在")
    p = (await db.execute(
        select(AdaptationProject).where(AdaptationProject.id == v.project_id)
    )).scalar_one()
    if p.user_id != current_user.id and not getattr(current_user, "is_superuser", False):
        raise HTTPException(404, "无权访问")
    return v


@router.post("/projects/{project_id}/runs", response_model=AdaptationVersionOut)
async def create_run(
    payload: RunCreate,
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    if not (p.metadata_ or {}).get("scene_boundaries"):
        raise HTTPException(400, "尚未切场，请先调用 /split")
    pipe = AdaptationPipeline(db=db, llm=get_default_service())
    version = await pipe.create_full_run(p, extra_prompt=payload.extra_prompt)
    asyncio.create_task(_background_run(p.id, version.id))
    return _version_to_out(version)


async def _background_run(project_id: int, version_id: int) -> None:
    """后台 task：用一个新的 db session 执行实际改写。"""
    from app.core.database import async_session
    async with async_session() as session:
        p = (await session.execute(
            select(AdaptationProject).where(AdaptationProject.id == project_id)
        )).scalar_one_or_none()
        v = (await session.execute(
            select(AdaptationVersion).where(AdaptationVersion.id == version_id)
        )).scalar_one_or_none()
        if not p or not v:
            return
        pipe = AdaptationPipeline(db=session, llm=get_default_service())
        try:
            await pipe.execute_full_run(p, v)
        except Exception as e:
            logger.exception("背景跑改编失败")
            v.status = "failed"
            v.error = str(e)[:500]
            await session.commit()
            await event_bus.publish(version_id, {"event": "version_failed", "error": str(e)[:200]})


@router.get("/projects/{project_id}/runs", response_model=List[AdaptationVersionOut])
async def list_runs(
    p: AdaptationProject = Depends(_get_owned_project),
    db: AsyncSession = Depends(get_db),
):
    rows = (await db.execute(
        select(AdaptationVersion).where(AdaptationVersion.project_id == p.id)
        .order_by(AdaptationVersion.version_no.desc())
    )).scalars().all()
    return [_version_to_out(v) for v in rows]


@router.get("/runs/{version_id}", response_model=VersionDetailOut)
async def get_version(
    v: AdaptationVersion = Depends(_get_owned_version),
    db: AsyncSession = Depends(get_db),
):
    from sqlalchemy.orm import selectinload
    # Re-query with eager loaded scene_results to avoid lazy load
    v_full = (await db.execute(
        select(AdaptationVersion)
        .where(AdaptationVersion.id == v.id)
        .options(selectinload(AdaptationVersion.scene_results))
    )).scalar_one()
    return {**_version_to_out(v_full), "scene_results": [_scene_to_out(s) for s in v_full.scene_results]}


@router.post("/runs/{version_id}/stream/ticket")
async def create_stream_ticket(
    v: AdaptationVersion = Depends(_get_owned_version),
    current_user: User = Depends(get_current_user),
):
    """为后续 EventSource 拉流颁发一次性短票据（5s 内有效）。

    浏览器原生 EventSource 无法附加 Authorization 头，因此把鉴权下沉到这一步：
    先用主 token POST 此接口拿 ticket，再用 ?ticket=… 拉 /stream。
    """
    return {"ticket": create_sse_ticket(current_user.id, v.id)}


@router.get("/runs/{version_id}/stream")
async def stream_run(version_id: int, ticket: str = Query(...)):
    if verify_sse_ticket(ticket, version_id) is None:
        raise HTTPException(401, "ticket 无效或已过期")
    sub = event_bus.subscribe(version_id)

    async def gen():
        try:
            yield "data: " + json.dumps({"event": "subscribed", "version_id": version_id}) + "\n\n"
            while True:
                try:
                    payload = await asyncio.wait_for(sub.queue.get(), timeout=15.0)
                    yield "data: " + json.dumps(payload, default=str) + "\n\n"
                    if payload.get("event") in ("version_done", "version_failed"):
                        return
                except asyncio.TimeoutError:
                    yield ": ping\n\n"
        finally:
            event_bus.unsubscribe(sub)

    return StreamingResponse(gen(), media_type="text/event-stream")


@router.post("/runs/{version_id}/scenes/{scene_index}/rerun", response_model=SceneResultOut)
async def rerun_scene(
    scene_index: int,
    payload: SceneRerunRequest,
    v: AdaptationVersion = Depends(_get_owned_version),
    db: AsyncSession = Depends(get_db),
):
    scene = (await db.execute(
        select(AdaptationSceneResult).where(
            AdaptationSceneResult.version_id == v.id,
            AdaptationSceneResult.scene_index == scene_index,
        )
    )).scalar_one_or_none()
    if not scene:
        raise HTTPException(404, "scene 不存在")
    if scene.status == "running":
        raise HTTPException(409, "该场正在跑，无法重跑")
    p = (await db.execute(
        select(AdaptationProject).where(AdaptationProject.id == v.project_id)
    )).scalar_one()
    pipe = AdaptationPipeline(db=db, llm=get_default_service())
    try:
        await pipe.rerun_scene(p, v, scene, payload.extra_prompt)
    except ValueError as e:
        raise HTTPException(409, str(e))
    return _scene_to_out(scene)


@router.patch("/runs/{version_id}/scenes/{scene_index}", response_model=SceneResultOut)
async def patch_scene(
    scene_index: int,
    payload: SceneManualPatch,
    v: AdaptationVersion = Depends(_get_owned_version),
    db: AsyncSession = Depends(get_db),
):
    scene = (await db.execute(
        select(AdaptationSceneResult).where(
            AdaptationSceneResult.version_id == v.id,
            AdaptationSceneResult.scene_index == scene_index,
        )
    )).scalar_one_or_none()
    if not scene:
        raise HTTPException(404, "scene 不存在")
    edits = list(scene.manual_edits or [])
    edits.append({
        "type": "manual", "at": utcnow_naive().isoformat(),
        "before": scene.rewritten_scene_text, "after": payload.rewritten_scene_text,
    })
    scene.manual_edits = edits
    scene.rewritten_scene_text = payload.rewritten_scene_text
    scene.status = "manual_edited"
    await db.commit()
    # 让 server-side onupdate 列重新水合，避免 _scene_to_out 触发同步 lazy load
    await db.refresh(scene)
    return _scene_to_out(scene)


@router.get("/runs/{version_id}/export")
async def export_run(
    format: str = "txt",
    v: AdaptationVersion = Depends(_get_owned_version),
    db: AsyncSession = Depends(get_db),
):
    if format not in ("txt", "docx"):
        raise HTTPException(400, "format 仅支持 txt/docx")
    scenes = (await db.execute(
        select(AdaptationSceneResult).where(AdaptationSceneResult.version_id == v.id)
        .order_by(AdaptationSceneResult.scene_index)
    )).scalars().all()

    def _ensure_header(text: str, title: Optional[str], scene_index: int) -> str:
        """确保改写结果以场号标题行+人物列表行开头；缺失则由代码兜底拼接。"""
        if not text:
            return text
        first_line = text.split("\n", 1)[0].strip()
        if any(p.search(first_line) for p in _SCENE_TITLE_PATTERNS):
            return text
        header = title or f"场{scene_index + 1}"
        char_names: list[str] = []
        for m in _CHAR_NAME_PATTERN.finditer(text):
            name = m.group(1)
            if name and name not in char_names:
                char_names.append(name)
        character_line = f"人物：{' '.join(char_names)}" if char_names else ""
        parts = [header]
        if character_line:
            parts.append(character_line)
        return "\n".join(parts) + "\n" + text

    parts = [
        _ensure_header(s.rewritten_scene_text or s.original_scene_text, s.scene_title, s.scene_index)
        for s in scenes
    ]

    if format == "txt":
        text = "\n\n".join(parts)
        return StreamingResponse(
            BytesIO(text.encode("utf-8")), media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=adaptation_v{v.version_no}.txt"},
        )

    # docx: 标题已在正文中，无需额外加 heading
    from docx import Document
    doc = Document()
    for s in scenes:
        content = _ensure_header(s.rewritten_scene_text or s.original_scene_text, s.scene_title, s.scene_index)
        for line in content.split("\n"):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if any(p.search(line_stripped) for p in _SCENE_TITLE_PATTERNS):
                doc.add_heading(line_stripped, level=2)
            elif line_stripped.startswith("人物："):
                p = doc.add_paragraph(line_stripped)
                for run in p.runs:
                    run.bold = True
            else:
                doc.add_paragraph(line_stripped)
    buf = BytesIO()
    doc.save(buf); buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=adaptation_v{v.version_no}.docx"},
    )
